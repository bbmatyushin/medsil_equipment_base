"""Создание документом для инженеров - актов"""

import os
import re
from pathlib import Path

from django.conf import settings
from django.db.models.query import QuerySet
from docx import Document
from docx.table import Table, _Cell
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Service


class CreateServiceAkt:
    def __init__(self, client: dict, job_content: str,
                 description: str, spare_parts: list,
                 template_path: Path,
                 accessories: QuerySet,
                 replacement_equipment: str,
                 accessories_with_quantity: list):
        self.akt = Document(str(template_path))
        self.file_name = template_path.name
        self.client = client
        self.description = description
        self.job_content = job_content
        self.spare_parts = spare_parts
        self.accessories = accessories
        self.replacement_equipment = replacement_equipment
        self.accessories_with_quantity = accessories_with_quantity
        self.save_file_path = self.create_save_path()

    def create_save_path(self) -> str:
        """Проверяет наличие папки для хранения актов определенной модели оборудования и
        возвращает путь расположения файла"""
        equipment_short_name = self.client['equipment_short_name'] \
            .replace(' ', '_') \
            .replace('/', '-')
        eq_dir_path = Path(settings.MEDIA_ROOT, 'docs', 'service_akt', equipment_short_name)
        save_file_name = (
            f"{self.file_name.split('.')[0]}_{self.client['{{ SERIAL_NUM }}'].replace(' ', '_')}"
            f"{'_' + self.client['{{ DATE }}'] if not self.client['{{ DATE }}'].startswith('___') else ''}"
            f".docx"
        )
        if not os.path.exists(eq_dir_path):
            os.mkdir(eq_dir_path)

        return str(Path(settings.MEDIA_ROOT, 'docs', 'service_akt',
                        equipment_short_name, save_file_name))

    def update_paragraphs(self):
        """Для обновления абзацев. НЕ ИСПОЛЬЗУЕТСЯ"""
        for par in self.akt.paragraphs:
            if 'PART_SERV' in par.text:
                par.text = par.text.replace('PART_SERV', 'NEW DETAIL')

    def update_tables(self):
        """Обноления данных в таблицах документа WORD"""
        # итерируемся по всем таблицам в документе
        for i, table in enumerate(self.akt.tables, start=1):
            if i == 1 and self.file_name in ["Akt_in_service.docx", "Akt_from_service.docx"]:
                self.head_table(table)

            if i == 2: self.main_table(table)  # таблица №2

            if i == 3: self.description_update(table)

            if i == 4:
                if self.file_name == "service_akt_MEDSIL.docx":
                    self.job_content_update(table)
                elif self.file_name in ["Akt_in_service.docx", "Akt_from_service.docx"] and len(self.spare_parts) > 0:
                    # self.fill_accessories_table_for_service(table)
                    self.spare_part_table(table, self.accessories_with_quantity)

            if i == 5:
                if self.file_name == "service_akt_MEDSIL.docx" and len(self.spare_parts) > 0:
                    self.spare_part_table(table, self.spare_parts)  # Используемые запчасти
                # Табличка с указанием подменного оборудования
                elif self.file_name == "Akt_in_service.docx" and self.replacement_equipment:
                    self.replacement_equipment_table(table)

            if i == 6 and self.file_name == "Akt_in_service.docx" \
                    and self.accessories:  # Табличка с перечнем коплектующих для подменного оборудования
                self.fill_accessories_table(table)

        self.akt.save(self.save_file_path)

    @staticmethod
    def cell_paragraph_gen(table: Table):
        """Возвращает абзац из ячейки таблицы"""
        for n, row in enumerate(table.rows, start=1):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield n, paragraph  # возвращает (номер строки, объект параграф)

    def head_table(self, table: Table):
        """Обновляем таблицу с указанием города и даты (сразу под названеим акта)"""
        for n, paragraph in self.cell_paragraph_gen(table):
            for key in self.client.keys():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, self.client[key])
                    if self.file_name == 'Akt_from_service.docx':
                        paragraph.runs[0].font.size = Pt(12)  # задаем разме шрифта
                    if key == "{{ AKT_DATE }}":
                        paragraph.paragraph_format.right_indent = Pt(0)  # отступ справа в 0 пунктов

    def main_table(self, table: Table):
        """Обновления даттых в таблице с реквизитами"""
        for n, paragraph in self.cell_paragraph_gen(table):
            for k, v in self.client.items():
                if k in paragraph.text:
                    paragraph.text = paragraph.text.replace(k, v)
                    if self.file_name == 'Akt_from_service.docx':
                        paragraph.runs[0].font.size = Pt(12)  # задаем разме шрифта
                    elif self.file_name == "Akt_in_service.docx":
                        paragraph.runs[0].font.size = Pt(11)  # задаем разме шрифта
                    if k == '{{ CLIENT }}':
                        paragraph.runs[0].bold = True
                    if k in ('{{ EQUIPMENT }}', '{{ SERIAL_NUM }}') \
                            and self.file_name == 'Akt_in_service.docx':
                        paragraph.runs[0].font.size = Pt(12)  # задаем разме шрифта

    def description_update(self, table: Table):
        """Обновление описания работ"""
        for n, paragraph in self.cell_paragraph_gen(table):
            if n == 4:
                paragraph.text = self.description

    def job_content_update(self, table):
        """Описание проведенных работ"""
        for n, paragraph in self.cell_paragraph_gen(table):
            if n == 1:
                paragraph.text = self.job_content

    def spare_part_table(self, table: Table, rows_data: list):
        """Обновляем данные в таблице Замененные детали.
        И комплектующих для актов приема-передачи в ремонт и из ремонта

        :param table - таблица с которой будем работать
        :param rows_data - список данных, которыми будет наполняться таблица
        """
        count_parts = len(rows_data)

        # Проходим по строкам, начиная со второй (индекс 1), т.к. первая — заголовки
        for row_idx in range(1, max(count_parts + 1, len(table.rows))):
            if row_idx >= len(table.rows):
                new_row = table.add_row()  # Добавляем новую строку, если не хватает
                # Добавляем границы для всех ячеек новой строки
                for cell in new_row.cells:
                    self._set_cell_borders(cell)
            else:
                new_row = table.rows[row_idx]

            # Заполняем ячейки строки, если есть соответствующая запчасть
            if row_idx <= count_parts and row_idx - 1 < len(rows_data):
                part = rows_data[row_idx - 1]  # Индекс запчасти: 0, 1, 2...
                cells = new_row.cells
                cells[0].text = str(row_idx)  # №
                if part[1]:
                    text_name = f"{part[0]} (арт. {part[1]})"  # Наименование + Артикул
                else:
                    text_name = f"{part[0]}"  # Наименование
                cells[1].text = text_name
                # Добавляем количество запчастей (предполагаем, что количество находится в part[2])
                # Если в spare_parts передаются кортежи (name, article, quantity)
                if len(part) >= 3:
                    cells[2].text = str(part[2])  # количество
                else:
                    cells[2].text = "1"  # значение по умолчанию, если количество не указано
                
                # Выравниваем количество по центру
                if len(cells) > 2:
                    self._set_cell_alignment(cells[2], align='center')

                # if self.file_name == 'Akt_from_service.docx':
                for cell in cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(11)  # задаем разме шрифта

    @staticmethod
    def _set_cell_alignment(cell: _Cell, align: str = 'center'):
        """Устанавливает выравнивание текста в ячейке
        
        Args:
            cell: ячейка таблицы
            align: тип выравнивания ('left', 'center', 'right', 'both', 'distribute')
        """
        # Получаем или создаем свойства ячейки
        cell_tcPr = cell._tc.get_or_add_tcPr()
        
        # Создаем элемент для вертикального выравнивания
        cell_alignment = OxmlElement('w:vAlign')
        # Преобразуем строковое значение в правильный формат для вертикального выравнивания
        if align == 'center':
            v_align = 'center'
        elif align == 'top':
            v_align = 'top'
        elif align == 'bottom':
            v_align = 'bottom'
        else:
            v_align = 'center'  # значение по умолчанию
            
        cell_alignment.set(qn('w:val'), v_align)
        
        # Добавляем выравнивание в свойства ячейки
        cell_tcPr.append(cell_alignment)
        
        # Также устанавливаем выравнивание для параграфа (горизонтальное выравнивание)
        # Импортируем необходимые константы
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        align_mapping = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE
        }
        
        paragraph_align = align_mapping.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        for paragraph in cell.paragraphs:
            paragraph.alignment = paragraph_align

    def fill_accessories_table(self, table: Table):
        """Заполняет таблицу комплектующих для подменного оборудования"""
        accessories_list = self.accessories.order_by("name")
        # Проходим по строкам, начиная со второй (индекс 1), если первая - заголовок
        for row_idx in range(1, max(len(accessories_list) + 1, len(table.rows))):
            if row_idx >= len(table.rows):
                new_row = table.add_row()  # Добавляем новую строку, если не хватает
                # Добавляем границы для всех ячеек новой строки
                for cell in new_row.cells:
                    self._set_cell_borders(cell)

            # Заполняем ячейки строки, если есть соответствующий аксессуар
            if row_idx - 1 < len(accessories_list):
                accessory = accessories_list[row_idx - 1]
                cells = table.rows[row_idx].cells
                cells[0].text = str(row_idx)  # №
                cells[1].text = accessory  # Наименование
                cells[2].text = "1"  # Количество по умолчанию
                # Выравниваем количество по центру
                if len(cells) > 2:
                    self._set_cell_alignment(cells[2], align='center')

                # Настройка шрифта
                for cell in cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(11)

    def replacement_equipment_table(self, table: Table):
        """Заполняет таблицу с подменным оборудованием для акта приема-передачи в ремонт"""
        # Таблица должна содержать одну строку и две ячейки
        if len(table.rows) > 0:
            row = table.rows[0]
            if len(row.cells) >= 2:
                # Заполняем вторую ячейку (индекс 1) значением подменного оборудования
                cell = row.cells[1]
                cell.text = self.replacement_equipment
                # Устанавливаем размер шрифта
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                # Устанавливаем границы ячейки
                self._set_cell_borders(cell)

    @staticmethod
    def _set_cell_borders(cell: _Cell):
        """Устанавливает границы вокруг ячеек"""
        cell_tcPr = cell._tc.get_or_add_tcPr()
        # Добавляет границы (рамку) вокруг ячейки вручную через OxmlElement
        tc_border = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tc_border.append(border)

        cell_tcPr.append(tc_border)


def create_service_atk(obj: Service, akt_name: str):
    """Создание акта для преданного объекта из модели Service.

    serviceAkt - Акт о проведении работ
    acceptInAkt - Акт приема-передачи оборудования в ремонт
    acceptFromAkt - Акт приема-передачи оборудования из ремонт
    """
    if akt_name == 'serviceAkt':
        template_name = 'service_akt_MEDSIL.docx'
    elif akt_name == 'acceptInAkt':
        template_name = 'Akt_in_service.docx'
    else:
        template_name = 'Akt_from_service.docx'

    template_path = Path(settings.MEDIA_ROOT, 'docs', 'service_akt', template_name)  # путь к шаблону
    dept = obj.equipment_accounting.equipment_acc_department_equipment_accounting.first().department
    client_city = dept.client.city.name if dept.client.city.name != 'Не указан' else ''
    address = f"{client_city} {dept.client.address if dept.client.address else ''}"
    
    # Получаем выбранное контактное лицо из сохраненных данных
    contact_person = "__________________________________"
    if hasattr(obj, 'contact_person_data') and obj.contact_person_data:
        contact_person_fio = obj.contact_person_data.get('fio', '').strip()
        contact_person_position = obj.contact_person_data.get('position', '').strip()
        
        # Формируем строку контактного лица с дополнительной информацией
        contact_person_parts = []
        if contact_person_fio:
            contact_person_parts.append(contact_person_fio)
        
        # Добавляем позицию, если она есть
        if contact_person_position:
            contact_person_parts.append(contact_person_position)
        
        # Получаем телефоны из объекта контактного лица
        contact_person_obj = None
        if hasattr(obj, 'contact_person') and obj.contact_person:
            contact_person_obj = obj.contact_person
        elif obj.contact_person_data.get('contact_person_id'):
            try:
                from ebase.models import DeptContactPers
                contact_person_obj = DeptContactPers.objects.get(
                    id=obj.contact_person_data['contact_person_id']
                )
            except (DeptContactPers.DoesNotExist, KeyError):
                pass
        
        # Добавляем телефоны в указанном порядке
        if contact_person_obj:
            if contact_person_obj.mob_phone:
                contact_person_parts.append(contact_person_obj.mob_phone)
            elif contact_person_obj.work_phone:
                contact_person_parts.append(contact_person_obj.work_phone)
        
        # Формируем итоговую строку
        if contact_person_parts:
            contact_person = ', '.join(contact_person_parts)
    
    # Определяем дату для поля {{ AKT_DATE }} в зависимости от типа акта
    # Словарь для перевода месяцев на русский
    month_translation = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    
    if akt_name == 'acceptInAkt':
        if obj.beg_dt:
            day = obj.beg_dt.day
            month = month_translation[obj.beg_dt.month]
            year = obj.beg_dt.year
            akt_date = f'«{day}» {month} {year} г.'
        else:
            akt_date = "«     »_______________202__г."
    elif akt_name == 'acceptFromAkt':
        if obj.end_dt:
            day = obj.end_dt.day
            month = month_translation[obj.end_dt.month]
            year = obj.end_dt.year
            akt_date = f'«{day}» {month} {year} г.'
        else:
            akt_date = "«     »_______________202__г."
    else:
        akt_date = "«     »_______________202__г."
    
    client = {
        '{{ CLIENT }}': dept.client.name,
        '{{ CLIENT_CONTACT }}': contact_person,
        '{{ ADDRESS }}': address,
        '{{ PHONE }}': dept.client.phone if dept.client.phone else '',
        '{{ EMAIL }}': dept.client.email if dept.client.email else '',
        '{{ INN }}': f"ИНН {dept.client.inn if dept.client.inn else ''}",
        '{{ KPP }}': f"КПП {dept.client.kpp}" if dept.client.kpp else 'КПП',
        '{{ EQUIPMENT }}': obj.equipment_accounting.equipment.full_name
        if obj.equipment_accounting.equipment.full_name else '',
        '{{ SERIAL_NUM }}': obj.equipment_accounting.serial_number,
        '{{ DATE }}': obj.end_dt.strftime('%d.%m.%Y') if obj.end_dt else '________________',
        '{{ AKT_DATE }}': akt_date,
        '{{ CITY }}': "__________________",
        'equipment_short_name': obj.equipment_accounting.equipment.short_name
        if obj.equipment_accounting.equipment.short_name else obj.equipment_accounting.equipment.full_name,
    }
    description = obj.description.replace("\r\n", "\n") if obj.description else ''
    job_content = obj.job_content.replace("\r\n", "\n") if obj.job_content else ''
    accessories = \
        obj.replacement_equipment.accessories.values_list("name", flat=True) if obj.replacement_equipment else []
    replacement_equipment = \
        f"{obj.replacement_equipment.equipment.full_name} (s/n {obj.replacement_equipment.serial_number})" \
            if obj.replacement_equipment else ""
    # Получаем информацию о запчастях с количеством
    spare_parts = []
    for spare_part in obj.spare_part.all():
        # Получаем количество из spare_part_count
        quantity = 1  # значение по умолчанию
        spare_part_id = str(spare_part.id)
        if spare_part_id in obj.spare_part_count:
            # Берем первое количество из списка
            part_info = obj.spare_part_count[spare_part_id]
            if isinstance(part_info, list) and len(part_info) > 0:
                quantity = part_info[0].get('service_part_count', 1)
            elif isinstance(part_info, dict):
                quantity = part_info.get('service_part_count', 1)
        spare_parts.append((spare_part.name, spare_part.article, quantity))
    
    # Получаем информацию о комплектующих с количеством из ServiceAccessories
    accessories_with_quantity = []
    for service_accessory in obj.service_accessories.all():
        accessory_name = service_accessory.accessory.name
        quantity = service_accessory.quantity
        accessories_with_quantity.append((accessory_name, '', quantity))

    #TODO: accessories можно создать таким же списком как и accessories_with_quantity,
    # чтобы использовать один метод для заполнения таблицы

    create_akt = CreateServiceAkt(client, job_content, description, spare_parts,
                                  template_path, accessories, replacement_equipment,
                                  accessories_with_quantity,)
    create_akt.update_tables()

    if akt_name == 'serviceAkt':
        obj.service_akt = create_akt.save_file_path
    elif akt_name == 'acceptInAkt':
        obj.accept_in_akt = create_akt.save_file_path
    else:
        obj.accept_from_akt = create_akt.save_file_path

    obj.save()

    return client['equipment_short_name'], client['{{ SERIAL_NUM }}']


if __name__ == '__main__':
    CreateServiceAkt(client={}).update_tables()
