"""Создание приказа о направлении работника в командировку."""

import os
from pathlib import Path

from django.conf import settings
from docx import Document
from docx.text.paragraph import Paragraph

from .models import BusinessTrip


def _replace_placeholders_in_paragraph(paragraph: Paragraph, replacements: dict):
    """Замена плейсхолдеров в абзаце с сохранением форматирования первого run.

    :param paragraph: абзац документа Word
    :param replacements: словарь {плейсхолдер: значение}
    """
    full_text = paragraph.text
    has_match = False
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            has_match = True
    if not has_match:
        return

    # Сохраняем форматирование первого run
    font_size = None
    bold = None
    font_name = None
    if paragraph.runs:
        source_run = paragraph.runs[0]
        font_size = source_run.font.size
        bold = source_run.font.bold
        font_name = source_run.font.name

    paragraph.clear()
    run = paragraph.add_run(full_text)
    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if font_name is not None:
        run.font.name = font_name


def _iter_all_paragraphs(doc):
    """Итерация по всем абзацам документа: основные + в ячейках таблиц."""
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def create_trip_order(obj: BusinessTrip):
    """Создание приказа о направлении работника в командировку.

    Шаблон: ``MEDIA_ROOT/docs/trips/order_trip.docx``
    Сохранение: ``MEDIA_ROOT/docs/trips/trip<doc_number>/Приказ_о_направлении_в_командировку_<doc_number>.docx``

    Плейсхолдеры шаблона:
        ``{{ DOC_NUM }}``       — ``BusinessTrip.doc_number``
        ``{{ CREATION_DATE}}``  — ``BusinessTrip.creation_date`` (DD.MM.YYYY)
        ``{{ FIO }}``           — Фамилия Имя Отчество сотрудника
        ``{{ DEPART_CITY }}``   — города подразделений через «; »
        ``{{ DEPARTMENT }}``    — наименования подразделений через «; »
        ``{{ BEG_DT }}``        — дата выезда (DD.MM.YYYY)
        ``{{ END_DT }}``        — дата возвращения (DD.MM.YYYY)
        ``{{ DAYS_COUNT }}``    — количество дней командировки

    :param obj: объект командировки
    :return: номер документа командировки
    """
    template_path = Path(settings.MEDIA_ROOT, "docs", "trips", "order_trip.docx")
    doc = Document(str(template_path))

    # --- ФИО сотрудника ---
    emp = obj.employee
    fio = " ".join(
        part for part in (emp.last_name, emp.first_name, emp.patron) if part
    ).strip()

    # --- Города и подразделения по пунктам: «город, подразделение» через «; » ---
    # Первые 2 пункта — на первую строку (P11), остальные — на P15.
    destinations = obj.destinations.select_related("department__city").all()
    dest_pairs: list[str] = []
    for dest in destinations:
        city_name = dest.city.name if dest.city else ""
        dept_name = dest.department.name if dest.department else ""
        pair = ", ".join(p for p in (city_name, dept_name) if p)
        if pair and pair not in dest_pairs:
            dest_pairs.append(pair)

    first_line = "; ".join(dest_pairs[:2])
    rest_line = "; ".join(dest_pairs[2:])

    # --- Словарь замен ---
    # Шаблон содержит «{{ DEPART_CITY }}, {{ DEPARTMENT }}» — рассчитан на один
    # пункт. Подставляем first_line в {{ DEPART_CITY }}, а {{ DEPARTMENT }}
    # очищаем, чтобы не дублировать.
    # Шаблон «{{ DEPART_CITY }}, {{ DEPARTMENT }}» рассчитан на один пункт.
    # Комбинированный ключ срабатывает первым — заменяет всю конструкцию целиком,
    # индивидуальные ключи ниже — fallback для нетипичного шаблона.
    replacements = {
        "{{ DEPART_CITY }}, {{ DEPARTMENT }}": first_line,
        "{{ DOC_NUM }}": str(obj.doc_number or ""),
        "{{ CREATION_DATE}}": (
            obj.creation_date.strftime("%d.%m.%Y") if obj.creation_date else ""
        ),
        "{{ FIO }}": fio,
        "{{ DEPART_CITY }}": first_line,
        "{{ DEPARTMENT }}": "",
        "{{ BEG_DT }}": obj.beg_dt.strftime("%d.%m.%Y") if obj.beg_dt else "",
        "{{ END_DT }}": obj.end_dt.strftime("%d.%m.%Y") if obj.end_dt else "",
        "{{ DAYS_COUNT }}": str(obj.days_count or ""),
    }

    # Запоминаем индекс абзаца с местом назначения — до замены
    depart_par_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "{{ DEPART_CITY }}" in paragraph.text:
            depart_par_idx = i
            break

    # Заменяем плейсхолдеры во всех абзацах (включая таблицы)
    for paragraph in _iter_all_paragraphs(doc):
        _replace_placeholders_in_paragraph(paragraph, replacements)

    # --- Перенос оставшихся пунктов на строку P15 ---
    # P12 — пустая, P13 — подпись «(место назначения…)», P14 — пустая, P15 — целевая.
    if rest_line and depart_par_idx is not None:
        next_idx = depart_par_idx + 4
        if next_idx < len(doc.paragraphs):
            overflow_par = doc.paragraphs[next_idx]
            overflow_par.clear()
            run = overflow_par.add_run(rest_line)
            # Копируем форматирование из абзаца-источника
            src_runs = doc.paragraphs[depart_par_idx].runs
            if src_runs:
                src = src_runs[0]
                if src.font.size:
                    run.font.size = src.font.size
                if src.font.name:
                    run.font.name = src.font.name

    # --- Сохранение ---
    trip_doc_number = f"trip{obj.doc_number}"
    save_dir = Path(settings.MEDIA_ROOT, "docs", "trips", trip_doc_number)
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)

    save_file_name = f"Приказ_о_направлении_в_командировку_{obj.doc_number}.docx"
    save_path = str(Path(save_dir, save_file_name))
    doc.save(save_path)

    obj.order_trip = save_path
    obj.save()

    return obj.doc_number
